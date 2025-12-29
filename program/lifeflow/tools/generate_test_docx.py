from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 输出路径
OUTPUT_PATH = r"c:\Users\29705\Desktop\plan4\test_docs\LifeFlow_Test_10pages.docx"

lorem = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Integer nec odio. Praesent libero. Sed cursus ante dapibus diam. "
    "Sed nisi. Nulla quis sem at nibh elementum imperdiet. Duis sagittis ipsum. "
    "Praesent mauris. Fusce nec tellus sed augue semper porta. Mauris massa. "
    "Vestibulum lacinia arcu eget nulla. Class aptent taciti sociosqu ad litora torquent per conubia nostra, per inceptos himenaeos. "
    "Curabitur sodales ligula in libero. Sed dignissim lacinia nunc."
)


def build_doc(pages: int = 10):
    doc = Document()

    # 标题
    title = doc.add_paragraph("LifeFlow 测试文档（10页）")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(20)

    # 元信息
    meta = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n用途: 文件解析/摘要提取测试")
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i in range(1, pages + 1):
        doc.add_paragraph("")
        h = doc.add_paragraph(f"第 {i} 页 - 测试段落")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(14)

        # 每页写入多个段落，确保分页
        for j in range(6):
            p = doc.add_paragraph(f"段落 {j+1}: {lorem} "*3)
            p_format = p.paragraph_format
            p_format.space_after = Pt(12)

        # 强制分页（除最后一页外）
        if i < pages:
            doc.add_page_break()

    return doc


def main():
    import os
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = build_doc(10)
    doc.save(OUTPUT_PATH)
    print(f"✅ 生成完成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
