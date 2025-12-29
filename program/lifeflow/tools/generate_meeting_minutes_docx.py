from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

OUTPUT_PATH = r"c:\Users\29705\Desktop\plan4\test_docs\LifeFlow_会议记录_10页.docx"

AGENDA = [
    "项目进度回顾",
    "需求变更评审",
    "技术方案讨论",
    "测试与上线计划",
    "风险与问题清单",
    "资源与排期调整",
    "行动项分配",
]

ATTENDEES = [
    "产品经理 - 张三",
    "技术负责人 - 李四",
    "后端工程师 - 王五",
    "前端工程师 - 赵六",
    "测试工程师 - 钱七",
    "运维 - 周八",
]

LOREM_CN = (
    "本次讨论围绕重点模块展开，针对接口性能、异常处理、数据一致性、边界条件、日志与监控进行了深入评估。"
    "我们确认了里程碑目标与交付物，明确了依赖关系与风险项，并制定了缓解策略与备选方案。"
)

DECISIONS = [
    "统一使用 DeepSeek 进行摘要与任务生成，关闭多 Provider 路由。",
    "数据库接口统一通过 Node 服务暴露，避免双写。",
    "文件解析限制为 PDF/DOCX/TXT，单文件不超过 100MB。",
    "前端任务导入必须用户确认，避免误导入。",
]

ACTIONS = [
    "后端：补充 API 文档与错误码说明（负责人：李四，截止：下周三）",
    "前端：实现 Vue 迁移 PoC 组件（负责人：赵六，截止：下周五）",
    "测试：补充文件解析与任务生成集成用例（负责人：钱七，截止：下周四）",
    "运维：完善启动脚本与日志采集（负责人：周八，截止：本周五）",
]


def add_heading(doc: Document, text: str, size: int = 16, center: bool = False):
    p = doc.add_paragraph(text)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(size)


def add_list(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(f"• {item}")
        p.paragraph_format.space_after = Pt(6)


def build_minutes(pages: int = 10):
    doc = Document()

    # 标题与元信息
    add_heading(doc, "LifeFlow 项目会议记录（10页）", size=20, center=True)
    meta = doc.add_paragraph(
        f"会议时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"会议主题：阶段性回顾与计划\n"
        f"会议地点：线上\n"
        f"记录人：系统自动生成"
    )
    meta.paragraph_format.space_after = Pt(12)

    # 参会人员
    add_heading(doc, "参会人员", size=14)
    add_list(doc, ATTENDEES)

    # 会议议程
    add_heading(doc, "会议议程", size=14)
    add_list(doc, AGENDA)

    # 主体内容分页写入
    for i in range(1, pages + 1):
        add_heading(doc, f"第 {i} 页 - 讨论要点", size=14)

        # 每页若干段落，包含讨论点/决定/行动项
        for j in range(1, 5):
            p = doc.add_paragraph(f"讨论点 {j}：{LOREM_CN}")
            p.paragraph_format.space_after = Pt(8)

        add_heading(doc, "会议决定", size=12)
        add_list(doc, DECISIONS)

        add_heading(doc, "行动项", size=12)
        add_list(doc, ACTIONS)

        if i < pages:
            doc.add_page_break()

    return doc


def main():
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = build_minutes(10)
    doc.save(OUTPUT_PATH)
    print(f"✅ 生成完成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
